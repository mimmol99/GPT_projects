import re
import time
import openai
import tkinter as tk
from tkinter import  simpledialog,messagebox
import requests
import os
import tiktoken
import math
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
import json
import sys
import nltk
from nltk.tokenize import sent_tokenize
from file_to_txt import file_to_text
from graphic_request import request_string,request_file_paths
from pptx import Presentation
from docx import Document
from tqdm import tqdm
from pydub import AudioSegment
from pydub.silence import split_on_silence
import magic



nltk.download('punkt')

# Get the absolute path to the parent directory
#parent_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Add it to the Python path
#sys.path.insert(0, parent_dir)




class GPTRequester:
    
    def __init__(self) -> None:
        
        self.base_path = os.getcwd()
        self.model_path = os.path.join(self.base_path,"model.txt")
        self.openaiapi_path = os.path.join(self.base_path,"openaiapi.txt")
        self.check_and_create_path(self.model_path)
        self.check_and_create_path(self.openaiapi_path)

        with open(self.openaiapi_path, 'r') as f:
            openai.api_key = f.read().strip()
    
        with open(self.model_path, 'r') as f:
            self.model = f.read().strip()

        self.models_limit_path = os.path.join(self.base_path,"models_limit.json")

        if not os.path.exists(self.models_limit_path):
            self.find_models_tokens_limit()

        self.previous_messages = []

        
    def request_gpt(self,model=None, request_phrase = "",text="", temperature=0.5):
        
        models = self.request_models()
    
        if model is None:
            model = self.model
            if model not in models:              
                model = self.select_item_from_list(models)
                self.model =  model
                with open(self.model_path, 'w') as f:
                    f.write(model)
            
        models_limit = {}
        
        if os.path.exists(self.models_limit_path):
            with open(self.models_limit_path, 'r') as json_file:
                models_limit = json.load(json_file)
                
            if len(models)!=len(models_limit.keys()):
                models_limit = self.find_models_tokens_limit()
        else:
            models_limit = self.find_models_tokens_limit()
        
        
        openai.Model.retrieve(model)   

        n_tokens = self.num_tokens_from_string(text) * 2 #to leave response space  
        tokens_limit = models_limit[model] //  2
        
        if n_tokens > tokens_limit:
            text_chunks = self.split_text(text,tokens_limit)
            total_response = ""
            for chunk in text_chunks:
                    try:
                        completion = openai.ChatCompletion.create(
                        model=model,
                        messages=[{"role": "user", "content":request_phrase+ chunk}],
                        temperature=temperature
                        )
                        response = completion["choices"][0]["message"]["content"]
                        
                        total_response = total_response + response
                        openai.Model.retrieve(model)
                    except:
                        time.sleep(10)
                        total_response = total_response + self.request_gpt(model=model, request_phrase = request_phrase,text=chunk, temperature=0.5)
                    
            return total_response
                
        else:
            try:
                completion = openai.ChatCompletion.create(
                model=model,
                messages=[{"role": "user", "content":request_phrase+text}],
                temperature=temperature
                )
                
                response = completion["choices"][0]["message"]["content"]
                
                openai.Model.retrieve(model)
            except:
                time.sleep(10)
                response = self.request_gpt(model=model, request_phrase = request_phrase,text=text, temperature=0.5)

            return response

    def request_gpt_translation(self,model="text-davinci-003",text="", temperature=0.5):
        min_tokens = 2048
        
        
        models = self.request_models()
    
        if model is None or model not in models:
            model = self.select_item_from_list(models)
            with open(self.model_path, 'w') as f:
                f.write(model)
            
        models_limit = {}
        
        if os.path.exists(self.models_limit_path):
            with open(self.models_limit_path, 'r') as json_file:
                models_limit = json.load(json_file)
                
            if len(models)!=len(models_limit.keys()):
                models_limit = self.find_models_tokens_limit()
        else:
            models_limit = self.find_models_tokens_limit()
        
        
        retrieve_response = openai.Model.retrieve(model)        
        n_tokens = self.num_tokens_from_string(text) * 2 #to leave response space  
        tokens_limit = models_limit[model]
        
        if n_tokens > tokens_limit:
            text_chunks = self.split_text(text,tokens_limit)
            total_response = ""
            for chunk in text_chunks:
                    completion =  openai.Completion.create(
    model="text-davinci-003",
    prompt="Translate this into Italian:\n\n "+text+" \n\n.",
    temperature=temperature,
    top_p=1.0,
    frequency_penalty=0.0,
    presence_penalty=0.0)
                    total_response = total_response + completion["choices"][0]["text"]
                    
            return total_response
                
        else:
            completion = openai.Completion.create(
    model="text-davinci-003",
    prompt="Translate this into Italian:\n "+text+".",
    temperature=temperature,
    top_p=1.0,
    frequency_penalty=0.0,
    presence_penalty=0.0)
            
            return completion["choices"][0]["text"]
        
    def request_models(self):
    
        self.check_and_create_path(self.openaiapi_path)
        
        with open(self.openaiapi_path, 'r') as f:
            openai.api_key = f.read().strip()# Define the headers for the request
            
        models = [model_dict["id"] for model_dict in openai.Model.list()["data"]]
        
        return models
   
    def check_and_create_path(self,path,askstring=" What text do you want to write to the file?"):
        if not os.path.exists(path):
            # create a root window
            root = tk.Tk()
            root.withdraw()  # hide the root window

            # ask the user if they want to create the path
            create_path = messagebox.askyesno(path+" does not exist", "Do you want to create the path?")
            if create_path:
                # get the text to write to the file
                text = simpledialog.askstring("Input",askstring)
                
                # create the path and write the text to it
                os.makedirs(os.path.dirname(path), exist_ok=True)
                with open(path, 'w') as f:
                    f.write(text)
            
            # destroy the root window
            root.destroy()
            
        
    def num_tokens_from_string(self,string, encoding_name=None):
        """Returns the number of tokens in a text string."""
        if encoding_name is None:
            encoding_name = "cl100k_base"
        encoding = tiktoken.get_encoding(encoding_name)
        num_tokens = len(encoding.encode(string))
        return num_tokens

    def split_text(self,text, max_tokens):
        # Tokenize the text into sentences
        sentences = sent_tokenize(text)

        chunks = []
        current_chunk = ''

        for sentence in sentences:
            # If adding the next sentence doesn't exceed the maximum token limit
            if len(current_chunk) + len(sentence) <= max_tokens:
                current_chunk += ' ' + sentence
            else:
                # Add the current chunk to the list
                chunks.append(current_chunk.strip())
                # Start a new chunk with the current sentence
                current_chunk = sentence
                
        # Don't forget to add the last chunk
        chunks.append(current_chunk.strip())

        return chunks
  

    def find_model_tokens_limit(self,model=None):
        
        # Read the model
        if model is None:
            with open(self.model_path, 'r') as f:
                model = f.read().strip()
                
        min_tokens = 2048
        
        
        # Setup Chrome Service
        webdriver_service = Service(ChromeDriverManager().install())
        
        # Create a new instance of the Chrome driver
        driver = webdriver.Chrome(service=webdriver_service)
        # Create a new instance of the Chrome driver
        driver = webdriver.Chrome(service=webdriver_service)
        
        url = "https://platform.openai.com/docs/models/" 

        # Go to the OpenAI documentation page
        driver.get(url)

        # Get all tables with the class 'models-table'
        models_tables = driver.find_elements(By.CLASS_NAME,"models-table")

        max_tokens = None
        
        # Process the models-table objects
        for table in models_tables:
            # Find all rows in the table, skipping the first one (header)

            rows = table.find_elements(By.CSS_SELECTOR,'tr')

            for row in rows:

                row_elements = row.find_elements(By.CSS_SELECTOR,"td")
                for re in row_elements:
                    current_model = row_elements[0].text
                    if current_model == model and "token" in re.text:
                        max_tokens = int(re.text.split(" ")[0].replace(",",""))
                        return max_tokens
                        
                        
        #close the driver
        driver.quit()
        
        return min_tokens

        
    def find_models_tokens_limit(self):
        
        models = self.request_models()
        
        min_tokens = 2048
        
        models_limit = {}
        
        # Setup Chrome Service
        webdriver_service = Service(ChromeDriverManager().install())
        
        # Create a new instance of the Chrome driver
        driver = webdriver.Chrome(service=webdriver_service)
        # Create a new instance of the Chrome driver
        driver = webdriver.Chrome(service=webdriver_service)
        
        url = "https://platform.openai.com/docs/models/" 

        # Go to the OpenAI documentation page
        driver.get(url)

        # Get all tables with the class 'models-table'
        models_tables = driver.find_elements(By.CLASS_NAME,"models-table")

        max_tokens = None
        
        # Process the models-table objects
        for table in models_tables:
            # Find all rows in the table, skipping the first one (header)

            rows = table.find_elements(By.CSS_SELECTOR,'tr')

            for row in rows:

                row_elements = row.find_elements(By.CSS_SELECTOR,"td")
                for re in row_elements:
                    current_model = row_elements[0].text
                    if current_model in models and "token" in re.text:
                        max_tokens = int(re.text.split(" ")[0].replace(",",""))
                        models_limit[current_model]= max_tokens 
                        #print(models_limit)
                        
        #close the driver
        driver.quit()
        
        for model in models:
            if model not in models_limit.keys():
                models_limit[model]= min_tokens  
        
                
        
        #print(models_limit)
        
        with open(self.models_limit_path, 'w') as json_file:
            json.dump(models_limit, json_file)

        return models_limit
    
    def select_item_from_list(self,item_list):
        # Create the root window
        root = tk.Tk()
        root.title('Select an item')

        # Create a StringVar() to store the selected item
        selected_item = tk.StringVar()

        # Create a Listbox widget
        listbox = tk.Listbox(root, exportselection=0, width=50,height=20)
        for item in item_list:
            listbox.insert(tk.END, item)
            listbox.pack()

        # Function to handle item selection
        def on_select(event):
        # Check if an item is selected
            if listbox.curselection():
                # Get selected item
                selection = event.widget.get(event.widget.curselection())
                selected_item.set(selection)
                # Show a message box with the selected item
                messagebox.showinfo("Selection", f"Selected: {selection}")
                # Close the window after selection
                root.withdraw()
            
            

        # Bind the select function to the listbox selection event
        listbox.bind('<<ListboxSelect>>', on_select)

        # Run the tkinter event loop
        root.mainloop()

        root.destroy()

        # Return the selected item
        return selected_item.get()


    def request_all_files(self,model=None,request_phrase="",document_name="output"):

        files = request_file_paths("select files")
        document = Document()
        requester = GPTRequester()
        pbar = tqdm(files, desc="Processing files")

        for f_path in pbar:
            base_name = os.path.basename(f_path).split(".")[0]
            pbar.set_description(f"Processing {base_name}")

            if self.is_audio_file(f_path) is False:
                text = file_to_text(f_path)
            else:
                text = self.transcript_audio(f_path)
            
            if text is None:continue
            response = requester.request_gpt(model = None,request_phrase=request_phrase,text=text)
            print(response)
            #response = text
            document.add_heading(base_name, level=1)
            document.add_paragraph(response)

        document.save(document_name+".docx")
        
    def transcript_audio(self,audio_path):
        audio_file= open(audio_path, "rb")
        transcript = openai.Audio.transcribe("whisper-1", audio_file)
        text = transcript["text"]
        return text
    
    
    def merge_audio_files(self,audio_paths):
        format_audio = os.path.basename(audio_paths[0]).split(".")[1]
        combined_audio = AudioSegment.from_file(audio_paths[0], format=format_audio)
            
        for audio in audio_paths[1:]:
            format_audio = os.path.basename(audio).split(".")[1]
            combined_audio+=AudioSegment.from_file(audio, format=format_audio)
            
        combined_audio.export("./combined_audio."+str(format_audio),format=format_audio)
        audio_path = "./combined_audio."+str(format_audio)
        return audio_path
    
    def is_audio_file(self,filepath):
        mime = magic.Magic(mime=True)
        mime_type = mime.from_file(filepath)
        return mime_type.startswith('audio/')
    
    
    
if __name__ == "__main__":
    #print(request_gpt_translation(model="ada",text="As of March 1, 2023, data sent to the OpenAI API will not be used to train or improve OpenAI models (unless you explitly opt in). One advantage to opting in is that the models may get better at your use case over time."))
    requester = GPTRequester()
    requester.request_all_files(request_phrase=request_string("Insert gpt prompt"),document_name=request_string("insert document name"))
