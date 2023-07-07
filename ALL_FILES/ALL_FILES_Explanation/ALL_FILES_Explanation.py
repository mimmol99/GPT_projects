import tqdm
import sys
import os
import PyPDF2
from docx import Document
from tqdm import tqdm
# Get the absolute path to the parent directory
parent_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Add it to the Python path
sys.path.insert(0, parent_dir)

from SHARED_FILES.directory_request import request_file_path,request_file_paths
from SHARED_FILES.divide_text_into_chunks import divide_text_into_chunks
from SHARED_FILES.request_gpt import request_gpt_translation
from PDF.PDFtoText import pdf_to_text
from PPT.PPT_explanation.PPT_explanation import explain_pptx,pptx_to_text
from TXT.TextExplanation.TextExplanation import text_explanation 


def main():
    files = request_file_paths("select files")
    document = Document()
    pbar = tqdm(files, desc="Processing files")
    for f_path in pbar:
        base_name = os.path.basename(f_path).split(".")[0]
        pbar.set_description(f"Processing {base_name}")
        
        if f_path.endswith(".txt"):
            total_text = text_explanation(path_text=f_path)
            #total_text = text_explanation(text=total_text,request_phrase="traduci in italiano: ")
            #total_text  = request_gpt_translation(model="ada",text=total_text)
        elif f_path.endswith(".pptx"):
            total_text = pptx_to_text(f_path)
            total_text = text_explanation(text=total_text)
            #total_text = text_explanation(text=total_text,request_phrase="traduci in italiano: ")
            #total_text  = request_gpt_translation(model="ada",text=total_text)
        elif f_path.endswith(".pdf"):
            total_text = pdf_to_text(f_path)
            total_text = text_explanation(text=total_text)
            #total_text = text_explanation(text=total_text,request_phrase="traduci in italiano: ")
            #total_text  = request_gpt_translation(model="ada",text=total_text)
        else:
            continue

        document.add_heading(os.path.basename(f_path).split(".")[0], level=1)
        document.add_paragraph(total_text)

    document.save("./out.docx")
        
if __name__ == "__main__":
    main()
